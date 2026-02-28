"""Generate project guide document as .docx"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Consolas'
code_style.font.size = Pt(9)
code_style.paragraph_format.space_before = Pt(4)
code_style.paragraph_format.space_after = Pt(4)
code_style.paragraph_format.left_indent = Cm(1)

def add_code_block(text):
    """Add a code block (multiple lines) to the document."""
    for line in text.strip().split('\n'):
        doc.add_paragraph(line, style='CodeBlock')

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    doc.add_paragraph()  # spacer

def add_func_table(functions):
    """Add a function reference table. functions = list of (name, params, returns, description)"""
    table = doc.add_table(rows=1 + len(functions), cols=4)
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(['Function / Method', 'Parameters', 'Returns', 'Description']):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for r_idx, (name, params, returns, desc) in enumerate(functions):
        table.rows[r_idx + 1].cells[0].text = name
        table.rows[r_idx + 1].cells[1].text = params
        table.rows[r_idx + 1].cells[2].text = returns
        table.rows[r_idx + 1].cells[3].text = desc
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Sign Language Recognition App', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('A Complete Project Guide for Students')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(16)
subtitle.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
doc.add_paragraph()
desc = doc.add_paragraph(
    'This document explains the architecture, data flow, and code structure of a real-time '
    'sign language recognition system built with MediaPipe and TensorFlow. '
    'It is designed for students who are new to machine learning and computer vision.'
)
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Introduction & Key Concepts',
    '2. Project Overview',
    '3. Setup & Installation',
    '4. Project File Structure',
    '5. How the Files Connect (Import Graph)',
    '6. The Processing Pipeline',
    '7. Module Reference',
    '   7.1 config.py',
    '   7.2 core/holistic_detector.py',
    '   7.3 core/landmark_extractor.py',
    '   7.4 core/frame_buffer.py',
    '   7.5 core/predictor.py',
    '   7.6 utils/preprocessing.py',
    '   7.7 utils/visualization.py',
    '   7.8 src/collect_data.py',
    '8. Data Collection Workflow',
    '9. Data Format & Storage',
    '10. Critical Rules (Don\'t Break These!)',
    '11. Project Phases & Roadmap',
]
for item in toc_items:
    doc.add_paragraph(item)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION & KEY CONCEPTS
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('1. Introduction & Key Concepts', level=1)

doc.add_paragraph(
    'Before diving into the code, here are the foundational concepts you need to understand. '
    'If you already know these, feel free to skip ahead to Section 2.'
)

doc.add_heading('What is MediaPipe?', level=2)
doc.add_paragraph(
    'MediaPipe is a free, open-source library by Google that can detect human body parts in images and video. '
    'In this project, we use its "Holistic" model which can detect hands, face, and body pose simultaneously. '
    'When MediaPipe sees a hand in a video frame, it identifies 21 key points (called "landmarks") on that hand \u2014 '
    'the wrist, each knuckle, each fingertip, etc. Each landmark has an (x, y, z) coordinate.'
)

doc.add_heading('What is a Feature Vector?', level=2)
doc.add_paragraph(
    'A feature vector is simply a list of numbers that represents something. In our case, we take the 21 landmarks '
    'from each hand (each with 3 coordinates: x, y, z), and flatten them into a single list of numbers. '
    'With 2 hands: 21 landmarks \u00d7 3 coordinates \u00d7 2 hands = 126 numbers. This 126-number list is our '
    '"feature vector" for one video frame.'
)

doc.add_heading('What is an LSTM?', level=2)
doc.add_paragraph(
    'LSTM (Long Short-Term Memory) is a type of neural network that is good at understanding sequences \u2014 '
    'data that has a time component. For example, the sign for "Hello" involves waving, which is a sequence of '
    'hand positions over time. Our LSTM model takes a sequence of 30 feature vectors (30 frames = 1 second of video) '
    'and predicts which sign is being performed.'
)

doc.add_heading('What is Normalization?', level=2)
doc.add_paragraph(
    'Normalization is the process of transforming data so that it is consistent regardless of external factors. '
    'For example, the same hand sign should produce similar numbers whether the hand is close to or far from the camera, '
    'or whether it is in the center or corner of the frame. We normalize by: (1) centering all landmarks on the wrist '
    '(so wrist is always at position 0,0,0), and (2) scaling by palm width (so hand size doesn\'t matter).'
)

doc.add_heading('What is OpenCV?', level=2)
doc.add_paragraph(
    'OpenCV (cv2 in Python) is a library for computer vision. In this project, we use it to: '
    'read video from the webcam, draw landmarks and UI elements on frames, display the video window, '
    'and handle keyboard input.'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 2. PROJECT OVERVIEW
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('2. Project Overview', level=1)
doc.add_paragraph(
    'This project is a real-time sign language recognition system. It uses a webcam to capture video of a person\'s '
    'hands, detects hand landmarks using MediaPipe, converts those landmarks into numerical feature vectors, '
    'and feeds sequences of those vectors into a trained LSTM model to predict which sign is being performed.'
)
doc.add_paragraph('Currently, the system recognizes 8 signs:')
add_table(
    ['Type', 'Signs'],
    [
        ['Static (alphabet)', 'A, B, C, D, E'],
        ['Dynamic (words)', 'Hello, ThankYou, Please'],
        ['Special', '_Neutral (resting / no sign)'],
    ]
)
doc.add_paragraph(
    'The project is divided into phases. Phase 1 (current) uses only hand landmarks (126 features). '
    'Phase 2 will add facial expressions (204 features). Phase 3 will add a real-time recognition app.'
)

# ═══════════════════════════════════════════════════════════════════════
# 3. SETUP
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('3. Setup & Installation', level=1)
doc.add_paragraph('Requirements: Python 3.10 or higher, and a webcam.')
doc.add_paragraph('Install all dependencies:')
add_code_block('pip install -r requirements.txt')
doc.add_paragraph('Key dependencies:')
add_table(
    ['Package', 'Purpose'],
    [
        ['numpy', 'Numerical arrays and math operations'],
        ['opencv-python', 'Webcam access, image processing, UI drawing'],
        ['mediapipe', 'Hand/face/pose landmark detection'],
        ['tensorflow-cpu', 'LSTM model loading and inference'],
        ['matplotlib', 'Plotting and visualization'],
        ['sounddevice', 'Audio feedback (beeps during recording)'],
        ['opencv-contrib-python', 'Extended OpenCV modules'],
    ]
)

# ═══════════════════════════════════════════════════════════════════════
# 4. FILE STRUCTURE
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('4. Project File Structure', level=1)
doc.add_paragraph('Here is the layout of the project and what each file does:')
add_code_block("""sign_language_app/
\u2502
\u251c\u2500\u2500 config.py                  # Central configuration (ALL settings live here)
\u2502
\u251c\u2500\u2500 core/                      # Core processing modules
\u2502   \u251c\u2500\u2500 __init__.py
\u2502   \u251c\u2500\u2500 holistic_detector.py   # MediaPipe wrapper (detects hands)
\u2502   \u251c\u2500\u2500 landmark_extractor.py  # Converts landmarks to feature vectors
\u2502   \u251c\u2500\u2500 frame_buffer.py        # Stores sequences of frames
\u2502   \u2514\u2500\u2500 predictor.py           # Loads model and makes predictions
\u2502
\u251c\u2500\u2500 utils/                     # Utility modules
\u2502   \u251c\u2500\u2500 __init__.py
\u2502   \u251c\u2500\u2500 preprocessing.py       # Normalization and data augmentation
\u2502   \u2514\u2500\u2500 visualization.py       # Drawing landmarks, UI panels, etc.
\u2502
\u251c\u2500\u2500 src/                       # User-facing applications
\u2502   \u251c\u2500\u2500 __init__.py
\u2502   \u2514\u2500\u2500 collect_data.py        # Data collection app (main application)
\u2502
\u251c\u2500\u2500 data/sequences/            # Training data (.npy files)
\u2502   \u251c\u2500\u2500 _Neutral/
\u2502   \u251c\u2500\u2500 A/ B/ C/ D/ E/
\u2502   \u251c\u2500\u2500 Hello/ ThankYou/ Please/
\u2502
\u251c\u2500\u2500 models/                    # Trained model files
\u2502   \u251c\u2500\u2500 best_model.keras       # The trained LSTM model
\u2502   \u2514\u2500\u2500 labels.pkl             # Label mapping (index -> sign name)
\u2502
\u251c\u2500\u2500 requirements.txt           # Python dependencies
\u2514\u2500\u2500 README.md                  # Project readme""")

# ═══════════════════════════════════════════════════════════════════════
# 5. IMPORT GRAPH
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('5. How the Files Connect (Import Graph)', level=1)
doc.add_paragraph(
    'Understanding which files depend on which is crucial. The arrows below show "imports from" relationships. '
    'config.py is at the center \u2014 every module imports settings from it.'
)
add_code_block("""                         config.py
                        (central hub)
                       /    |    |    \\
                      /     |    |     \\
                     v      v    v      v
       holistic_detector  landmark_extractor  frame_buffer  predictor
                             |
                             v
                        preprocessing.py

       visualization.py  (imports config + mediapipe drawing utils)

       collect_data.py  (imports ALL of the above except predictor)
         \u251c\u2500\u2500> holistic_detector
         \u251c\u2500\u2500> landmark_extractor
         \u251c\u2500\u2500> frame_buffer
         \u2514\u2500\u2500> visualization""")

doc.add_paragraph(
    'Key insight: config.py has NO project imports (only standard library "os"). It is the root dependency. '
    'collect_data.py is the "leaf" that pulls everything together.'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 6. PROCESSING PIPELINE
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('6. The Processing Pipeline', level=1)
doc.add_paragraph(
    'This is the most important diagram in this document. It shows how data flows from the webcam all the way '
    'to a prediction. Every module in the project serves a role in this pipeline.'
)
add_code_block("""  WEBCAM (captures video frame)
     |
     v
  HolisticDetector.detect(frame)
     |  Uses MediaPipe to find hand landmarks
     |  Fixes left/right hand swap for mirrored video
     v
  LandmarkExtractor.extract_and_normalize(results)
     |  Converts MediaPipe landmarks to numpy arrays
     |  Calls preprocessing.normalize_hand_landmarks()
     |    Step 1: Center all points on the wrist (wrist = origin)
     |    Step 2: Scale by palm width (hand size doesn't matter)
     |  Flattens to 126-dim feature vector
     v
  FrameBuffer.add(features)
     |  Stores the feature vector in a circular buffer
     |  Buffer holds up to 60 frames (2 seconds)
     v
  FrameBuffer.get_last_n(30)
     |  Extracts last 30 frames as a (30, 126) array
     v
  Predictor.predict(sequence)
     |  Feeds (30, 126) sequence into LSTM model
     |  Returns: predicted sign name + confidence score
     v
  OUTPUT: "Hello" (95.2% confidence)""")

doc.add_paragraph(
    'During data collection, the pipeline stops at the FrameBuffer stage \u2014 instead of predicting, '
    'the 30-frame sequence is saved as a .npy file for training.'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 7. MODULE REFERENCE
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('7. Module Reference', level=1)
doc.add_paragraph(
    'This section covers every module in detail, including all public functions, their parameters, '
    'and what they return. Each module can be tested individually by running it directly.'
)

# ── 7.1 config.py ──
doc.add_heading('7.1 config.py \u2014 Central Configuration', level=2)
doc.add_paragraph(
    'This file holds ALL configurable settings for the entire project. Instead of having "magic numbers" '
    'scattered across files, everything is defined here in a single Config class. You never create an instance \u2014 '
    'just access settings directly like Config.SEQUENCE_LENGTH.'
)
doc.add_paragraph('Usage in any file:')
add_code_block("""from config import Config
print(Config.SEQUENCE_LENGTH)  # 30
print(Config.MIN_CONFIDENCE)   # 0.7""")

doc.add_heading('Configuration Groups', level=3)
add_table(
    ['Group', 'Key Settings', 'Default Values'],
    [
        ['Paths', 'DATA_DIR, MODEL_DIR, MODEL_PATH, LABELS_PATH', 'data/sequences, models/, models/unified_model.tflite, models/labels.pkl'],
        ['Landmarks', 'HAND_LANDMARKS, COORDS_PER_LANDMARK, FEATURES_PER_FRAME', '21, 3, 126'],
        ['Body Parts', 'USE_HANDS, USE_FACE, USE_POSE', 'True, False, False'],
        ['Sequences', 'SEQUENCE_LENGTH, BUFFER_SIZE', '30 frames, 60 frames'],
        ['Detection', 'MIN_DETECTION_CONFIDENCE, MIN_TRACKING_CONFIDENCE', '0.7, 0.5'],
        ['Prediction', 'MIN_CONFIDENCE, NEUTRAL_LABEL, IGNORE_CLASSES', '0.7, "_Neutral", ["_Neutral"]'],
        ['Stability', 'STABILITY_THRESHOLD, STABILITY_FRAMES, COOLDOWN_FRAMES', '0.02, 10, 15'],
        ['Camera', 'CAMERA_INDEX, FRAME_WIDTH, FRAME_HEIGHT, FPS', '0, 640, 480, 30'],
        ['Data Collection', 'SIGNS_TO_COLLECT, SAMPLES_PER_SIGN, SAMPLES_FOR_NEUTRAL', '8 signs, 50, 100'],
        ['UI', 'SHOW_LANDMARKS, SHOW_DEBUG_INFO, colors (BGR)', 'True, True, green/white/red'],
    ]
)

doc.add_heading('Methods', level=3)
add_func_table([
    ('get_features_per_frame()', 'None', 'int', 'Calculates total features based on USE_HANDS and USE_FACE flags. Returns 126 (hands only) or 204 (hands + face).'),
    ('get_data_path(sign_name)', 'sign_name: str', 'str', 'Returns full folder path for a sign\'s data, e.g. "data/sequences/A".'),
    ('ensure_directories()', 'None', 'None', 'Creates data/ and models/ directories and a subfolder for each sign. Call at startup.'),
])

doc.add_heading('Landmark Index Map', level=3)
doc.add_paragraph('MediaPipe hand landmarks are numbered 0-20. Key indices used in normalization:')
add_code_block("""        8   12  16  20      <- Fingertips
        |   |   |   |
        7   11  15  19
        |   |   |   |
        6   10  14  18
        |   |   |   |
        5   9   13  17      <- Knuckles (MCP joints)
         \\  |   |  /
          \\ |   | /
       4   \\|   |/
       |    \\   /
       3     \\ /
       |      1             <- THUMB_CMC (index 1)
       2      |
        \\     |
         \\    |
          \\   |
           \\  |
            \\ |
             \\|
              0             <- WRIST (index 0)

Key indices:
  WRIST_INDEX = 0           (center point for normalization)
  THUMB_CMC_INDEX = 1       (thumb base)
  INDEX_MCP_INDEX = 5       (index finger knuckle)
  PINKY_MCP_INDEX = 17      (pinky knuckle)

Palm width = distance from index 5 to index 17""")

doc.add_paragraph('Test: Run python config.py to verify all settings print correctly.')
doc.add_page_break()

# ── 7.2 holistic_detector.py ──
doc.add_heading('7.2 core/holistic_detector.py \u2014 MediaPipe Wrapper', level=2)
doc.add_paragraph(
    'This module wraps Google\'s MediaPipe Holistic model into an easy-to-use class. '
    'Its main job is to take a video frame (image) and return the detected hand landmarks. '
    'It also fixes an important bug: when using a mirrored webcam view (which is natural for users), '
    'MediaPipe swaps left and right hands. This module corrects that swap automatically.'
)

doc.add_heading('Class: HolisticDetector', level=3)
add_func_table([
    ('__init__(min_detection_confidence, min_tracking_confidence, static_image_mode, flip_hands)',
     'All optional. Defaults from Config. flip_hands=True enables hand swap fix.',
     'None',
     'Creates MediaPipe Holistic detector with given confidence thresholds.'),
    ('detect(frame)',
     'frame: BGR numpy array (from OpenCV)',
     'MediaPipe results object',
     'Converts BGR to RGB, runs detection, optionally swaps hands. Returns results with .left_hand_landmarks and .right_hand_landmarks.'),
    ('_swap_hands(results)',
     'results: MediaPipe results',
     'results (modified)',
     'Internal method. Swaps left_hand_landmarks with right_hand_landmarks to correct mirrored view.'),
    ('hands_detected(results)',
     'results: MediaPipe results',
     'bool',
     'Returns True if at least one hand is detected.'),
    ('get_hand_count(results)',
     'results: MediaPipe results',
     'int (0, 1, or 2)',
     'Returns how many hands are detected.'),
    ('get_detection_info(results)',
     'results: MediaPipe results',
     'dict',
     'Returns {"left_hand": bool, "right_hand": bool, "face": bool, "pose": bool}.'),
    ('release()',
     'None',
     'None',
     'Cleans up MediaPipe resources. Also works via context manager (with statement).'),
])

doc.add_heading('Why flip_hands?', level=3)
doc.add_paragraph(
    'When you look at a webcam, the image is mirrored (like a mirror). Your right hand appears on the left '
    'side of the screen. MediaPipe labels hands based on their screen position, so it calls your right hand '
    '"left" and vice versa. flip_hands=True swaps them back so the labels match reality.'
)

doc.add_paragraph('Test: Run python core/holistic_detector.py to open webcam and see hand detection in action.')
doc.add_page_break()

# ── 7.3 landmark_extractor.py ──
doc.add_heading('7.3 core/landmark_extractor.py \u2014 Feature Extraction', level=2)
doc.add_paragraph(
    'This is the bridge between MediaPipe\'s raw detection results and the numerical feature vectors that the '
    'LSTM model needs. It takes MediaPipe results, extracts the hand landmarks, normalizes them using the '
    'preprocessing module, and produces a 126-dimensional feature vector.'
)

doc.add_heading('Class: LandmarkExtractor', level=3)
add_func_table([
    ('__init__(config)',
     'config: Config class (optional, uses default)',
     'None',
     'Stores landmark indices and feature dimensions from config.'),
    ('extract_and_normalize(results)',
     'results: MediaPipe results from HolisticDetector',
     'numpy array, shape (126,)',
     'Main method. Extracts landmarks from both hands, normalizes each hand, flattens and concatenates into a single feature vector. Missing hands are filled with zeros.'),
    ('get_raw_landmarks(results)',
     'results: MediaPipe results',
     'tuple of (left, right) numpy arrays, each (21, 3) or None',
     'Returns un-normalized landmark arrays for each hand.'),
    ('get_wrist_position(results)',
     'results: MediaPipe results',
     'dict {"left": (x,y,z) or None, "right": (x,y,z) or None}',
     'Returns the wrist (x, y, z) position for each detected hand.'),
    ('hands_detected(results)',
     'results: MediaPipe results',
     'bool',
     'Returns True if at least one hand is detected.'),
    ('get_hand_count(results)',
     'results: MediaPipe results',
     'int',
     'Returns number of detected hands (0, 1, or 2).'),
    ('get_palm_size(results)',
     'results: MediaPipe results',
     'dict {"left": float or None, "right": float or None}',
     'Returns palm width (distance from index knuckle to pinky knuckle) for each hand.'),
])

doc.add_heading('Internal Data Flow', level=3)
add_code_block("""MediaPipe results
  |
  +--> Check: is left hand present? is right hand present?
  |
  +--> For each present hand:
  |      mediapipe_to_numpy(hand_landmarks) -> (21, 3) array
  |
  +--> landmarks_to_feature_vector(left, right)
  |      For each hand:
  |        normalize_hand_landmarks() -> centered & scaled (21, 3)
  |      Flatten each to (63,)
  |      Concatenate: left(63) + right(63) = (126,)
  |
  +--> Return (126,) feature vector""")

doc.add_paragraph('Test: Run python core/landmark_extractor.py to see live feature extraction from webcam.')
doc.add_page_break()

# ── 7.4 frame_buffer.py ──
doc.add_heading('7.4 core/frame_buffer.py \u2014 Circular Buffer', level=2)
doc.add_paragraph(
    'The LSTM model needs a sequence of frames (not just one frame) to make a prediction. '
    'The FrameBuffer stores incoming feature vectors in a circular buffer. "Circular" means that when the buffer '
    'is full, the oldest frame is removed to make room for the newest one. Think of it like a conveyor belt: '
    'new frames go in one end, old frames fall off the other.'
)

doc.add_heading('Class: FrameBuffer', level=3)
add_func_table([
    ('__init__(size, features_per_frame)',
     'size: int (default 60), features_per_frame: int (default 126)',
     'None',
     'Creates an empty buffer that holds up to "size" frames.'),
    ('add(frame)',
     'frame: numpy array, shape (126,)',
     'None',
     'Adds a feature vector to the buffer. If full, removes the oldest frame first.'),
    ('get_last_n(n)',
     'n: int (e.g. 30)',
     'numpy array, shape (n, 126)',
     'Returns the most recent n frames as a 2D array. This is what you feed to the LSTM model.'),
    ('get_all()',
     'None',
     'numpy array, shape (current_size, 126)',
     'Returns all frames currently in the buffer.'),
    ('get_motion_score(n_frames)',
     'n_frames: int (default 10)',
     'float',
     'Calculates average movement between consecutive recent frames. Used to detect if the hand has stopped moving.'),
    ('is_ready(required_frames)',
     'required_frames: int (default 30)',
     'bool',
     'Returns True if the buffer has at least "required_frames" stored.'),
    ('clear()',
     'None',
     'None',
     'Removes all frames from the buffer.'),
    ('__len__()',
     'None',
     'int',
     'Returns current number of frames in buffer.'),
    ('get_stats()',
     'None',
     'dict',
     'Returns buffer statistics: current size, max size, total frames added, whether ready.'),
])

doc.add_heading('How the Circular Buffer Works', level=3)
add_code_block("""Buffer size = 5 (simplified example)

add(frame1): [frame1]
add(frame2): [frame1, frame2]
add(frame3): [frame1, frame2, frame3]
add(frame4): [frame1, frame2, frame3, frame4]
add(frame5): [frame1, frame2, frame3, frame4, frame5]  <- FULL
add(frame6): [frame2, frame3, frame4, frame5, frame6]  <- frame1 removed
add(frame7): [frame3, frame4, frame5, frame6, frame7]  <- frame2 removed

get_last_n(3) -> [frame5, frame6, frame7]  <- most recent 3""")

doc.add_paragraph('Test: Run python core/frame_buffer.py to see buffer behavior with test data.')
doc.add_page_break()

# ── 7.5 predictor.py ──
doc.add_heading('7.5 core/predictor.py \u2014 Model Inference', level=2)
doc.add_paragraph(
    'This module loads a pre-trained Keras LSTM model and uses it to classify sign language gestures. '
    'It takes a sequence of 30 frames (shape: 30 rows \u00d7 126 columns) and outputs the predicted sign name '
    'along with a confidence score.'
)
doc.add_paragraph(
    'Important: The model files (best_model.keras and labels.pkl) are NOT included in the repository. '
    'They must be obtained separately (e.g., from Google Drive after training).'
)

doc.add_heading('Class: Predictor', level=3)
add_func_table([
    ('__init__(model_path, labels_path)',
     'model_path: str (default "models/best_model.keras"), labels_path: str (default from Config)',
     'None',
     'Loads the Keras model and label mapping. Raises FileNotFoundError if files are missing.'),
    ('predict(sequence)',
     'sequence: numpy array, shape (30, 126)',
     'tuple (class_name: str, confidence: float, probabilities: array)',
     'Runs inference. Returns predicted sign name, confidence (0-1), and probability for each class.'),
    ('predict_top_k(sequence, k)',
     'sequence: (30, 126) array, k: int (default 3)',
     'list of (class_name, confidence) tuples',
     'Returns the top k predictions sorted by confidence.'),
    ('is_neutral(predicted_class)',
     'predicted_class: str',
     'bool',
     'Returns True if the prediction is "_Neutral" (not a sign).'),
    ('get_class_names()',
     'None',
     'list of str',
     'Returns list of all sign names the model knows.'),
    ('get_num_classes()',
     'None',
     'int',
     'Returns number of sign classes.'),
    ('get_model_info()',
     'None',
     'dict',
     'Returns dict with sequence_length, features_per_frame, num_classes, class_names.'),
])

doc.add_heading('How labels.pkl Works', level=3)
doc.add_paragraph(
    'The labels file is a Python dictionary (pickled) that maps integer indices to sign names. '
    'For example: {0: "_Neutral", 1: "A", 2: "B", ...}. When the model outputs probabilities, '
    'the index with the highest probability is looked up in this dictionary to get the sign name.'
)

doc.add_paragraph('Test: Run python core/predictor.py (requires model files) to test loading and inference speed.')
doc.add_page_break()

# ── 7.6 preprocessing.py ──
doc.add_heading('7.6 utils/preprocessing.py \u2014 Normalization & Augmentation', level=2)
doc.add_paragraph(
    'This is arguably the most critical module in the project. It contains the normalization functions '
    'that transform raw MediaPipe landmarks into consistent, comparable feature vectors. '
    'THE SAME normalization must be used everywhere: data collection, training, and inference. '
    'If you change anything here, ALL training data must be recollected and the model must be retrained.'
)

doc.add_heading('Functions', level=3)
add_func_table([
    ('normalize_hand_landmarks(landmarks, wrist_idx, index_mcp_idx, pinky_mcp_idx)',
     'landmarks: (21, 3) array, wrist_idx: int, index_mcp_idx: int, pinky_mcp_idx: int',
     'numpy array, shape (21, 3)',
     'THE CORE FUNCTION. Step 1: Center on wrist (subtract wrist position). Step 2: Scale by palm width (divide by distance between index and pinky knuckles). Uses MIN_SCALE (0.001) to prevent division by zero.'),
    ('calculate_distance(point1, point2)',
     'point1: (3,) array, point2: (3,) array',
     'float',
     'Euclidean distance in 3D space: sqrt((x2-x1)\u00b2 + (y2-y1)\u00b2 + (z2-z1)\u00b2).'),
    ('landmarks_to_feature_vector(left_hand, right_hand)',
     'left_hand: (21, 3) array or None, right_hand: (21, 3) array or None',
     'numpy array, shape (126,)',
     'Normalizes each hand independently, flattens to (63,), concatenates left+right. Missing hands become zeros.'),
    ('mediapipe_to_numpy(hand_landmarks)',
     'hand_landmarks: MediaPipe landmark object',
     'numpy array, shape (21, 3)',
     'Converts MediaPipe\'s landmark format to a plain numpy array of (x, y, z) coordinates.'),
    ('augment_sequence(sequence, noise_level)',
     'sequence: (30, 126) array, noise_level: float (default 0.01)',
     'numpy array, shape (30, 126)',
     'Adds random Gaussian noise for data augmentation during training.'),
    ('mirror_sequence(sequence)',
     'sequence: (30, 126) array',
     'numpy array, shape (30, 126)',
     'Swaps left/right hands and negates x-coordinates. Creates horizontally flipped training data.'),
    ('scale_sequence(sequence, scale_factor)',
     'sequence: (30, 126) array, scale_factor: float (default random 0.9-1.1)',
     'numpy array, shape (30, 126)',
     'Scales all coordinates uniformly. Simulates slightly different hand sizes.'),
])

doc.add_heading('The 3-Step Normalization Process (Detailed)', level=3)
add_code_block("""STEP 0: MediaPipe gives landmarks in 0-1 range
  (relative to image dimensions)

STEP 1: CENTER ON WRIST
  wrist_position = landmarks[0]     # e.g. (0.5, 0.6, 0.02)
  centered = landmarks - wrist_position
  # Now wrist is at (0, 0, 0)
  # All other points are relative to wrist

STEP 2: SCALE BY PALM WIDTH
  palm_width = distance(landmarks[5], landmarks[17])
  # landmarks[5] = index finger knuckle
  # landmarks[17] = pinky knuckle
  if palm_width < 0.001:
      palm_width = 0.001    # prevent division by zero
  normalized = centered / palm_width
  # Now coordinates are in "palm width units"
  # Hand close to camera = same values as hand far away""")

doc.add_paragraph(
    'Why is this important? Without normalization, the same sign would produce very different numbers '
    'depending on where the hand is in the frame and how far it is from the camera. Normalization makes '
    'the feature vectors consistent, which is essential for the LSTM model to learn correctly.'
)

doc.add_paragraph('Test: Run python utils/preprocessing.py to test normalization and augmentation.')
doc.add_page_break()

# ── 7.7 visualization.py ──
doc.add_heading('7.7 utils/visualization.py \u2014 Drawing & UI', level=2)
doc.add_paragraph(
    'This module contains all the OpenCV drawing functions used to create the user interface overlay '
    'on the webcam video. Every function takes a frame (image) and draws on it in-place (modifies the original). '
    'Colors are in BGR format (Blue, Green, Red) which is OpenCV\'s default.'
)

doc.add_heading('Functions', level=3)
add_func_table([
    ('draw_landmarks(frame, results, show_hands, show_face)',
     'frame: BGR image, results: MediaPipe results, show_hands: bool, show_face: bool',
     'None (modifies frame)',
     'Draws the MediaPipe skeleton (dots at landmarks, lines connecting them) on the frame using green color.'),
    ('draw_hand_box(frame, results)',
     'frame: BGR image, results: MediaPipe results',
     'None (modifies frame)',
     'Draws bounding rectangles around detected hands with padding.'),
    ('draw_info_panel(frame, info_dict, position)',
     'frame: BGR image, info_dict: dict, position: "top" or "bottom"',
     'None (modifies frame)',
     'Draws a semi-transparent panel showing key-value pairs (e.g., "Status: Recording").'),
    ('draw_progress_bar(frame, progress, position, label, width, height, color)',
     'frame: BGR image, progress: float 0-1, other params optional',
     'None (modifies frame)',
     'Draws a progress bar with percentage text. Used during recording to show how many frames collected.'),
    ('draw_recording_indicator(frame, is_recording)',
     'frame: BGR image, is_recording: bool',
     'None (modifies frame)',
     'Shows red dot + "REC" when recording, gray "IDLE" when not. Top-right corner.'),
    ('draw_countdown(frame, seconds_remaining)',
     'frame: BGR image, seconds_remaining: int',
     'None (modifies frame)',
     'Draws a large centered countdown number (3... 2... 1...) with semi-transparent overlay.'),
    ('draw_instructions(frame, instructions, position)',
     'frame: BGR image, instructions: list of str, position: str',
     'None (modifies frame)',
     'Draws a panel with instruction lines at the bottom of the frame.'),
    ('draw_sign_label(frame, sign_name, samples_collected, samples_target)',
     'frame: BGR image, sign_name: str, samples: int, target: int',
     'None (modifies frame)',
     'Green-bordered panel at top center showing "Sign: X" and "Samples: n/target".'),
    ('draw_message(frame, message, message_type)',
     'frame: BGR image, message: str, type: "info"/"success"/"warning"/"error"',
     'None (modifies frame)',
     'Colored message at center-bottom. Color depends on type: yellow/green/orange/red.'),
])

doc.add_paragraph('Test: Run python utils/visualization.py to see all drawing functions on live webcam.')
doc.add_page_break()

# ── 7.8 collect_data.py ──
doc.add_heading('7.8 src/collect_data.py \u2014 Data Collection App', level=2)
doc.add_paragraph(
    'This is the main user-facing application. It opens a webcam window and guides you through collecting '
    'training data for each sign. It uses a state machine to manage the recording workflow.'
)

doc.add_heading('Class: DataCollector', level=3)
doc.add_paragraph('Key methods:')
add_func_table([
    ('__init__()',
     'None',
     'None',
     'Creates directories, initializes detector/extractor/buffer, opens webcam, sets initial state.'),
    ('run()',
     'None',
     'None',
     'Main loop: reads frames, processes, updates state, draws UI, handles keyboard input. Call this to start the app.'),
    ('_process_frame(frame)',
     'frame: BGR image',
     'tuple (results, features)',
     'Detects hands and extracts features from one frame.'),
    ('_update_recording(features)',
     'features: (126,) array',
     'None',
     'Handles the countdown/recording state machine. Adds frames to buffer when recording.'),
    ('_start_countdown()',
     'None',
     'None',
     'Begins 3-second countdown before recording starts.'),
    ('_start_recording()',
     'None',
     'None',
     'Transitions from countdown to active recording. Clears buffer.'),
    ('_stop_recording(save)',
     'save: bool',
     'None',
     'Ends recording. If save=True, saves buffer as .npy file. If False, discards.'),
    ('_delete_last_sample()',
     'None',
     'None',
     'Deletes the most recently saved .npy file for the current sign (retry/undo).'),
    ('_next_sign() / _prev_sign()',
     'None',
     'None',
     'Navigate through the list of signs to collect.'),
    ('_jump_to_sign(index)',
     'index: int',
     'None',
     'Jump directly to a sign by number (0-9 keys).'),
    ('_draw_ui(frame, results)',
     'frame: BGR image, results: MediaPipe results',
     'None',
     'Draws all UI elements: landmarks, sign label, instructions, recording indicator, progress bar, messages.'),
    ('_handle_input(key)',
     'key: int (OpenCV keycode)',
     'bool',
     'Processes keyboard input. Returns False if user pressed Q (quit).'),
    ('_count_existing_samples()',
     'None',
     'None',
     'Counts .npy files in each sign\'s folder to track progress.'),
    ('_get_next_filename(sign)',
     'sign: str',
     'str',
     'Generates path like "data/sequences/A/sequence_003.npy" based on existing files.'),
])
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 8. DATA COLLECTION WORKFLOW
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('8. Data Collection Workflow', level=1)
doc.add_paragraph('Run the data collection app:')
add_code_block('python -m src.collect_data')

doc.add_heading('Keyboard Controls', level=2)
add_table(
    ['Key', 'Action'],
    [
        ['SPACE', 'Start countdown / cancel recording'],
        ['N', 'Next sign'],
        ['P', 'Previous sign'],
        ['R', 'Retry (delete last saved sample)'],
        ['0-9', 'Jump to sign by number'],
        ['Q', 'Quit'],
    ]
)

doc.add_heading('State Machine', level=2)
doc.add_paragraph(
    'The recording process follows a strict state machine. Understanding this is key to understanding '
    'how collect_data.py works:'
)
add_code_block("""State Machine:

  IDLE  (waiting for user)
    |
    | [User presses SPACE]
    v
  COUNTDOWN  (3... 2... 1...)
    |
    | [3 seconds pass]
    v
  RECORDING  (capturing 30 frames)
    |
    +---> [30 frames captured] --> SAVE --> back to IDLE
    |
    +---> [User presses SPACE] --> CANCEL --> back to IDLE

During RECORDING:
  - Each frame's features are added to the buffer
  - If hand tracking is lost, zeros are added instead
  - Progress bar shows frames collected / 30
  - After 30 frames, sequence is saved as .npy file""")

doc.add_heading('What Gets Saved', level=2)
doc.add_paragraph(
    'Each recording saves a single .npy file containing a numpy array of shape (30, 126). '
    'The file is saved to data/sequences/{sign_name}/sequence_NNN.npy where NNN is an auto-incrementing number.'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 9. DATA FORMAT
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('9. Data Format & Storage', level=1)

doc.add_heading('File Structure', level=2)
add_code_block("""data/sequences/
\u251c\u2500\u2500 _Neutral/
\u2502   \u251c\u2500\u2500 sequence_001.npy    # shape: (30, 126)
\u2502   \u251c\u2500\u2500 sequence_002.npy
\u2502   \u2514\u2500\u2500 ... (target: 100 samples)
\u251c\u2500\u2500 A/
\u2502   \u251c\u2500\u2500 sequence_001.npy
\u2502   \u2514\u2500\u2500 ... (target: 50 samples)
\u251c\u2500\u2500 B/
\u251c\u2500\u2500 C/
\u251c\u2500\u2500 D/
\u251c\u2500\u2500 E/
\u251c\u2500\u2500 Hello/
\u251c\u2500\u2500 ThankYou/
\u2514\u2500\u2500 Please/""")

doc.add_heading('Array Layout', level=2)
doc.add_paragraph('Each .npy file contains a 2D numpy array:')
add_table(
    ['Dimension', 'Size', 'Meaning'],
    [
        ['Rows', '30', '30 frames at 30 FPS = 1 second of video'],
        ['Columns', '126', '126 features per frame'],
    ]
)
doc.add_paragraph('The 126 columns are organized as:')
add_table(
    ['Indices', 'Content'],
    [
        ['0-62', 'Left hand: 21 landmarks \u00d7 3 coordinates (x,y,z), flattened'],
        ['63-125', 'Right hand: 21 landmarks \u00d7 3 coordinates (x,y,z), flattened'],
    ]
)
doc.add_paragraph(
    'If a hand is not detected in a frame, its 63 values are all zeros. '
    'The model learns to handle zero-filled hands.'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 10. CRITICAL RULES
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('10. Critical Rules (Don\'t Break These!)', level=1)

doc.add_heading('Rule 1: Normalization Consistency', level=2)
doc.add_paragraph(
    'The normalize_hand_landmarks() function in utils/preprocessing.py MUST produce identical results '
    'during data collection, training, and inference. If you change the normalization logic (even slightly), '
    'you must: (1) delete ALL existing training data, (2) recollect all data, and (3) retrain the model. '
    'There is no shortcut.'
)

doc.add_heading('Rule 2: Feature Vector Layout', level=2)
doc.add_paragraph(
    'The order of features in the 126-dim vector must never change. Left hand is always indices 0-62, '
    'right hand is always 63-125. The LSTM model was trained on this exact layout.'
)

doc.add_heading('Rule 3: Sequence Length', level=2)
doc.add_paragraph(
    'The model expects exactly 30 frames. Sending more or fewer frames will cause errors or incorrect predictions. '
    'If you change SEQUENCE_LENGTH, you must retrain the model.'
)

doc.add_heading('Rule 4: Config is the Single Source of Truth', level=2)
doc.add_paragraph(
    'Never hardcode values that exist in config.py. If you need to change a threshold, camera resolution, '
    'or any other setting, change it in config.py. All modules read from Config.'
)

doc.add_heading('Rule 5: FEATURES_PER_FRAME Must Be Updated Manually', level=2)
doc.add_paragraph(
    'Config has two ways to get the feature count: Config.FEATURES_PER_FRAME (a static number, defaults to 126) '
    'and Config.get_features_per_frame() (calculated from USE_HANDS/USE_FACE). If you enable USE_FACE, '
    'you must also manually update FEATURES_PER_FRAME from 126 to 204.'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 11. PHASES
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading('11. Project Phases & Roadmap', level=1)

add_table(
    ['Phase', 'Status', 'Features', 'Details'],
    [
        ['Phase 1', 'Current', '126 features (hands only)', '8 signs: _Neutral, A-E, Hello, ThankYou, Please. Uses LSTM model with 30-frame sequences.'],
        ['Phase 2', 'Planned', '204 features (hands + face)', 'Adds 26 facial landmarks (78 features) for expression-based signs. Set USE_FACE=True in config.'],
        ['Phase 3', 'Planned', 'Real-time app', 'recognize.py \u2014 continuous real-time sign recognition using webcam with gesture detection and prediction filtering.'],
    ]
)

doc.add_heading('What Changes in Phase 2', level=2)
doc.add_paragraph(
    'Phase 2 adds facial expression recognition. The key changes will be: '
    '(1) Set USE_FACE = True in config.py and update FEATURES_PER_FRAME to 204, '
    '(2) Update landmark_extractor.py to extract 26 face landmarks, '
    '(3) Recollect all training data with the new 204-feature format, '
    '(4) Retrain the LSTM model with input shape (30, 204).'
)

# ═══════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Project_Guide.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
